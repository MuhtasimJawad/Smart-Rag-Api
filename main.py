import os
import io
import re
import json
import base64
import logging
import tempfile
import sqlite3
import threading
import unicodedata
from math import sqrt
from typing import List, Dict, Any, Optional, Tuple, Iterable
from uuid import uuid4

import numpy as np
import requests
import faiss
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import pandas as pd
from docx import Document
from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from starlette.concurrency import run_in_threadpool
from sentence_transformers import SentenceTransformer

# LangChain + HF
from langchain_huggingface import HuggingFaceEndpoint, ChatHuggingFace
from langchain_core.messages import SystemMessage, HumanMessage
from huggingface_hub.utils import HfHubHTTPError
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Optional: load .env if present (before reading env vars)
try:
    from dotenv import load_dotenv, find_dotenv
    load_dotenv(find_dotenv(usecwd=True), override=False)
    load_dotenv(os.path.join(os.path.dirname(__file__), ".env"), override=False)
except Exception:
    pass

# -------------------------
# Config & Logging
# -------------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("rag-api")

STORAGE_DIR = os.getenv("STORAGE_DIR", "storage")
os.makedirs(STORAGE_DIR, exist_ok=True)

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")

# LLM via LangChain + HF Endpoint
HUGGINGFACEHUB_API_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN")  # required for gated models
LLM_REPO_ID = os.getenv("LLM_REPO_ID", "meta-llama/Llama-3.1-8B-Instruct")
LLM_TEMPERATURE = float(os.getenv("LLM_TEMPERATURE", "0.2"))
LLM_MAX_NEW_TOKENS = int(os.getenv("LLM_MAX_NEW_TOKENS", "384"))
LLM_TOP_P = float(os.getenv("LLM_TOP_P", "0.95"))
LLM_TOP_K = int(os.getenv("LLM_TOP_K", "50"))
LLM_TIMEOUT = int(os.getenv("LLM_TIMEOUT", "60"))  # may be used by newer libs

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "900"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))
TOP_K_DEFAULT = int(os.getenv("TOP_K", "5"))
MAX_DB_ROWS = int(os.getenv("MAX_DB_ROWS", "200"))
PDF_OCR_DPI = int(os.getenv("PDF_OCR_DPI", "200"))
PDF_OCR_MAX_PIXELS = int(os.getenv("PDF_OCR_MAX_PIXELS", "5000000"))  # cap raster size for OCR

TESSERACT_PATH = os.getenv("TESSERACT_PATH")  # e.g., C:/Program Files/Tesseract-OCR/tesseract.exe
if TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

# -------------------------
# Icons / Labels for response metadata
# -------------------------
FILE_TYPE_ICON = {
    "pdf": "📕",
    "docx": "📄",
    "txt": "📝",
    "image": "🖼️",
    "csv": "🧮",
    "sqlite": "🗃️",
    "unknown": "📦",
}
FILE_TYPE_LABEL = {
    "pdf": "PDF document",
    "docx": "Word document",
    "txt": "Text file",
    "image": "Image",
    "csv": "CSV file",
    "sqlite": "SQLite database",
    "unknown": "File",
}


def file_icon_and_label(meta: Dict[str, Any]) -> Tuple[str, str]:
    t = (meta.get("type") or "unknown").lower()
    icon = FILE_TYPE_ICON.get(t, FILE_TYPE_ICON["unknown"])
    label = FILE_TYPE_LABEL.get(t, FILE_TYPE_LABEL["unknown"])
    return icon, label


# -------------------------
# Text cleaning & chunking
# -------------------------
def clean_text(text: str) -> str:
    """
    Normalize and lightly clean extracted text while preserving paragraph structure.
    - Normalize Unicode (NFKC)
    - Normalize newlines to \n
    - Remove zero-width and BOM chars
    - De-hyphenate word breaks across line endings (e.g., 'trans-\nform' -> 'transform')
    - Collapse repeated spaces/tabs within lines
    - Collapse 3+ blank lines to 2 newlines
    """
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
    text = re.sub(r"[\u200B-\u200D\uFEFF]", "", text)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        line = re.sub(r"[ ]{2,}", " ", line)
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def get_text_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""],
    )


def chunk_text(text: str) -> List[str]:
    text = clean_text(text)
    if not text:
        return []
    splitter = get_text_splitter()
    return splitter.split_text(text)


# -------------------------
# Helpers
# -------------------------
def normalize(vecs: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(vecs, axis=1, keepdims=True) + 1e-12
    return vecs / norms


def image_to_text(img: Image.Image) -> str:
    try:
        text = pytesseract.image_to_string(img)
        return clean_text(text)
    except Exception:
        logger.exception("OCR failed")
        return ""


def decode_base64_image(b64: str) -> Image.Image:
    try:
        img_bytes = base64.b64decode(b64)
        return Image.open(io.BytesIO(img_bytes)).convert("RGB")
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 image data")


def download_to_temp(url: str) -> str:
    r = requests.get(url, stream=True, timeout=60)
    if r.status_code != 200:
        raise HTTPException(status_code=400, detail=f"Failed to download: {r.status_code}")
    suffix = ""
    if "content-disposition" in r.headers:
        cd = r.headers["content-disposition"]
        m = re.search(r'filename="?([^"]+)"?', cd)
        if m:
            suffix = os.path.splitext(m.group(1))[1]
    if not suffix:
        suffix = os.path.splitext(url.split("?")[0].split("/")[-1])[1]
    fd, tmp_path = tempfile.mkstemp(suffix=suffix or "")
    with os.fdopen(fd, "wb") as f:
        for chunk in r.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    return tmp_path


# -------------------------
# Parsers
# -------------------------
def render_page_pixmap_for_ocr(page: fitz.Page) -> Image.Image:
    w_pt = page.rect.width
    h_pt = page.rect.height
    scale_dpi = PDF_OCR_DPI / 72.0
    width_px = w_pt * scale_dpi
    height_px = h_pt * scale_dpi
    pixels = width_px * height_px
    if pixels > PDF_OCR_MAX_PIXELS and pixels > 0:
        from math import sqrt
        factor = sqrt(PDF_OCR_MAX_PIXELS / pixels)
        scale = scale_dpi * factor
    else:
        scale = scale_dpi
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img


def parse_pdf_streaming(file_path: str, filename: str) -> Iterable[Tuple[str, Dict[str, Any]]]:
    doc = fitz.open(file_path)
    try:
        for i, page in enumerate(doc):
            page_num = i + 1
            try:
                text = page.get_text("text", sort=True) or ""
            except Exception:
                text = ""
            text = clean_text(text)

            if len(text) < 30:
                try:
                    img = render_page_pixmap_for_ocr(page)
                    text = image_to_text(img)
                except Exception as e:
                    logger.warning(f"OCR on page {page_num} failed: {e}")
                    text = ""

            if not text:
                continue

            for chunk_index, ch in enumerate(chunk_text(text)):
                meta = {"filename": filename, "type": "pdf", "page": page_num, "chunk_index": chunk_index}
                yield ch, meta
    finally:
        doc.close()


def parse_docx(file_path: str, filename: str) -> List[Tuple[str, Dict[str, Any]]]:
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for t in doc.tables:
        for row in t.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(row_text))
    full_text = clean_text("\n\n".join(parts))
    chunks = chunk_text(full_text)
    return [(ch, {"filename": filename, "type": "docx", "page": None, "chunk_index": i}) for i, ch in enumerate(chunks)]


def parse_txt(file_path: str, filename: str) -> List[Tuple[str, Dict[str, Any]]]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = clean_text(f.read())
    chunks = chunk_text(text)
    return [(ch, {"filename": filename, "type": "txt", "page": None, "chunk_index": i}) for i, ch in enumerate(chunks)]


def parse_image_file(file_path: str, filename: str) -> List[Tuple[str, Dict[str, Any]]]:
    img = Image.open(file_path).convert("RGB")
    text = image_to_text(img)
    chunks = chunk_text(text) if text else []
    return [(ch, {"filename": filename, "type": "image", "page": None, "chunk_index": i}) for i, ch in enumerate(chunks)]


def parse_csv(file_path: str, filename: str) -> List[Tuple[str, Dict[str, Any]]]:
    try:
        df = pd.read_csv(file_path)
    except Exception:
        df = pd.read_csv(file_path, sep=None, engine="python")
    text = clean_text(df.to_csv(index=False))
    chunks = chunk_text(text)
    return [(ch, {"filename": filename, "type": "csv", "page": None, "chunk_index": i}) for i, ch in enumerate(chunks)]


def parse_sqlite(file_path: str, filename: str) -> List[Tuple[str, Dict[str, Any]]]:
    con = sqlite3.connect(file_path)
    cur = con.cursor()
    results = []
    try:
        tables = cur.execute("SELECT name, type FROM sqlite_master WHERE type IN ('table','view')").fetchall()
        for tname, ttype in tables:
            schema_rows = cur.execute(f"PRAGMA table_info('{tname}')").fetchall()
            schema_lines = []
            for row in schema_rows:
                cid, name, ctype, notnull, dflt, pk = row
                schema_lines.append(f"- {name} {ctype} NOT NULL:{bool(notnull)} PK:{bool(pk)} DEFAULT:{dflt}")
            schema_text = "Schema:\n" + "\n".join(schema_lines) if schema_lines else "Schema: (unknown)"
            rows = []
            try:
                rows = cur.execute(f"SELECT * FROM '{tname}' LIMIT {MAX_DB_ROWS}").fetchall()
                colnames = [d[1] for d in cur.execute(f"PRAGMA table_info('{tname}')").fetchall()]
            except Exception:
                colnames = []
            content_lines = []
            if colnames:
                content_lines.append("Columns: " + ", ".join(colnames))
            for r in rows:
                try:
                    content_lines.append("Row: " + ", ".join([str(x) for x in r]))
                except Exception:
                    content_lines.append("Row: <unprintable>")
            table_text = f"SQLite {ttype}: {tname}\n{schema_text}\n" + "\n".join(content_lines)
            for i, ch in enumerate(chunk_text(table_text)):
                meta = {"filename": filename, "type": "sqlite", "table": tname, "page": None, "chunk_index": i}
                results.append((ch, meta))
    finally:
        cur.close()
        con.close()
    return results


def parse_any(file_path: str, original_filename: Optional[str] = None) -> List[Tuple[str, Dict[str, Any]]]:
    filename = original_filename or os.path.basename(file_path)
    ext = os.path.splitext(filename.lower())[1]
    if ext in [".pdf"]:
        return list(parse_pdf_streaming(file_path, filename))
    elif ext in [".docx"]:
        return parse_docx(file_path, filename)
    elif ext in [".txt", ".md", ".log"]:
        return parse_txt(file_path, filename)
    elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"]:
        return parse_image_file(file_path, filename)
    elif ext in [".csv"]:
        return parse_csv(file_path, filename)
    elif ext in [".db", ".sqlite", ".sqlite3"]:
        return parse_sqlite(file_path, filename)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


# -------------------------
# Vector Store with FAISS
# -------------------------
class VectorStore:
    def __init__(self, storage_dir: str, embedder: SentenceTransformer):
        self.storage_dir = storage_dir
        self.index_path = os.path.join(storage_dir, "index.faiss")
        self.meta_path = os.path.join(storage_dir, "metadata.jsonl")
        self.embedder = embedder
        self.dim = embedder.get_sentence_embedding_dimension()
        self.lock = threading.Lock()
        self.index = None
        self._load()

    def _load(self):
        if os.path.exists(self.index_path):
            self.index = faiss.read_index(self.index_path)
            logger.info(f"FAISS index loaded: {self.index_path}")
        else:
            self.index = faiss.IndexFlatIP(self.dim)
            logger.info(f"FAISS index initialized (dim={self.dim})")
        if not os.path.exists(self.meta_path):
            with open(self.meta_path, "w", encoding="utf-8"):
                pass

    def add_texts(self, texts: List[str], metadatas: List[Dict[str, Any]]) -> List[int]:
        assert len(texts) == len(metadatas)
        with self.lock:
            embeddings = self.embedder.encode(texts, batch_size=64, convert_to_numpy=True, show_progress_bar=False)
            embeddings = normalize(np.array(embeddings).astype("float32"))
            self.index.add(embeddings)
            ids = []
            with open(self.meta_path, "a", encoding="utf-8") as f:
                start_id = self.index.ntotal - len(texts)
                for i, (text, meta) in enumerate(zip(texts, metadatas)):
                    vec_id = start_id + i
                    meta_record = dict(meta)
                    meta_record["id"] = int(vec_id)
                    meta_record["text"] = text
                    f.write(json.dumps(meta_record, ensure_ascii=False) + "\n")
                    ids.append(vec_id)
            faiss.write_index(self.index, self.index_path)
            return ids

    def _iter_meta(self) -> List[Dict[str, Any]]:
        items = []
        with open(self.meta_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    items.append(json.loads(line))
                except Exception:
                    continue
        return items

    def search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        if self.index.ntotal == 0:
            return []
        q_emb = self.embedder.encode([query], convert_to_numpy=True, show_progress_bar=False)
        q_emb = normalize(q_emb.astype("float32"))
        D, I = self.index.search(q_emb, top_k)
        ids = I[0].tolist()
        scores = D[0].tolist()
        meta_items = self._iter_meta()
        meta_by_id = {m["id"]: m for m in meta_items}
        results = []
        for idx, score in zip(ids, scores):
            if idx == -1:
                continue
            m = meta_by_id.get(idx)
            if not m:
                continue
            res = dict(m)
            res["score"] = float(score)
            results.append(res)
        return results


# -------------------------
# LLM via LangChain + Hugging Face Endpoint
# -------------------------
class LLMClient:
    def __init__(
        self,
        repo_id: str = LLM_REPO_ID,
        token: Optional[str] = HUGGINGFACEHUB_API_TOKEN,
        temperature: float = LLM_TEMPERATURE,
        max_new_tokens: int = LLM_MAX_NEW_TOKENS,
        top_p: float = LLM_TOP_P,
        top_k: int = LLM_TOP_K,
        timeout: int = LLM_TIMEOUT,  # may be supported in newer versions
    ):
        self.repo_id = repo_id
        self.token = token

        self.endpoint = HuggingFaceEndpoint(
            repo_id=repo_id,
            huggingfacehub_api_token=token,
            temperature=temperature,
            max_new_tokens=max_new_tokens,
            top_p=top_p,
            top_k=top_k,
            do_sample=temperature > 0,
            return_full_text=False,
        )
        self.chat = ChatHuggingFace(llm=self.endpoint)

    def generate(self, messages: List[Any]) -> str:
        try:
            result = self.chat.invoke(messages)
            return (result.content or "").strip()
        except HfHubHTTPError as e:
            status = getattr(getattr(e, "response", None), "status_code", None)
            if status in (401, 403):
                raise HTTPException(
                    status_code=status,
                    detail="Hugging Face auth/permissions error. Check HUGGINGFACEHUB_API_TOKEN and model access (accept the model license).",
                )
            if status == 503:
                raise HTTPException(status_code=503, detail="Model is loading/warming up on Hugging Face. Try again shortly.")
            raise HTTPException(status_code=502, detail=f"Hugging Face inference error: {str(e)}")
        except Exception as e:
            logger.exception("LLM generation failed")
            raise HTTPException(status_code=500, detail=f"LLM generation failed: {e}")


# -------------------------
# RAG Prompt (chat messages)
# -------------------------
def build_messages(question: str, contexts: List[Dict[str, Any]], max_chars: int = 8000, per_ctx_limit: int = 2000) -> List[Any]:
    system_prompt = (
        "You are a careful, concise assistant. Use ONLY the provided context to answer. "
        "If the answer cannot be determined from the context, say: I don't know."
    )

    ctx_lines = []
    used = 0
    for i, c in enumerate(contexts, 1):
        src = c.get("filename", "unknown")
        page = c.get("page")
        where = f"{src}" + (f", page {page}" if page else "")
        text = (c.get("text") or "")[:per_ctx_limit]
        block = f"[{i}] Source: {where}\n{text}\n"
        if used + len(block) > max_chars:
            break
        ctx_lines.append(block)
        used += len(block)

    human_prompt = f"Context:\n{''.join(ctx_lines)}\nQuestion: {question}\nAnswer:"
    return [SystemMessage(content=system_prompt), HumanMessage(content=human_prompt)]


# -------------------------
# FastAPI app and State
# -------------------------
app = FastAPI(title="Smart RAG API", version="1.6.0")

class QueryRequest(BaseModel):
    question: str
    image_base64: Optional[str] = None
    top_k: Optional[int] = TOP_K_DEFAULT


@app.on_event("startup")
def startup_event():
    logger.info("Loading embedding model...")
    embedder = SentenceTransformer(EMBEDDING_MODEL)
    app.state.embedder = embedder
    app.state.vstore = VectorStore(STORAGE_DIR, embedder)
    app.state.llm = LLMClient(
        repo_id=LLM_REPO_ID,
        token=HUGGINGFACEHUB_API_TOKEN,
        temperature=LLM_TEMPERATURE,
        max_new_tokens=LLM_MAX_NEW_TOKENS,
        top_p=LLM_TOP_P,
        top_k=LLM_TOP_K,
        timeout=LLM_TIMEOUT,
    )
    logger.info("Startup complete.")


@app.get("/health")
def health():
    return {
        "status": "ok",
        "index_size": app.state.vstore.index.ntotal,
        "embedder": EMBEDDING_MODEL,
        "llm_repo": LLM_REPO_ID,
        "pdf_ocr_dpi": PDF_OCR_DPI,
        "pdf_ocr_max_pixels": PDF_OCR_MAX_PIXELS,
    }


# -------------------------
# Memory-safe ingestion helpers
# -------------------------
def ingest_pdf_streaming_to_store(file_path: str, filename: str, vstore: VectorStore, batch: int = 64, file_id: Optional[str] = None) -> int:
    count = 0
    texts, metas = [], []
    for text, meta in parse_pdf_streaming(file_path, filename):
        if file_id:
            meta["file_id"] = file_id
        texts.append(text)
        metas.append(meta)
        if len(texts) >= batch:
            vstore.add_texts(texts, metas)
            count += len(texts)
            texts, metas = [], []
    if texts:
        vstore.add_texts(texts, metas)
        count += len(texts)
    return count


# -------------------------
# Upload endpoint (returns file_id)
# -------------------------
@app.post("/upload")
async def upload(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded.")

    total_chunks = 0
    file_summaries = []

    async def process_one(file: UploadFile):
        suffix = os.path.splitext(file.filename)[1].lower()
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        file_id = f"file_{uuid4().hex[:12]}"
        try:
            with os.fdopen(fd, "wb") as f:
                content = await file.read()
                f.write(content)

            if suffix == ".pdf":
                chunks_count = await run_in_threadpool(
                    ingest_pdf_streaming_to_store, tmp_path, file.filename, app.state.vstore, 64, file_id
                )
                return {"filename": file.filename, "file_id": file_id, "chunks": chunks_count}
            else:
                parsed = await run_in_threadpool(parse_any, tmp_path, file.filename)
                if not parsed:
                    return {"filename": file.filename, "file_id": file_id, "chunks": 0}
                texts = [t for t, _ in parsed]
                metas = [m for _, m in parsed]
                # Attach file_id to each chunk
                for m in metas:
                    m["file_id"] = file_id
                ids = await run_in_threadpool(app.state.vstore.add_texts, texts, metas)
                return {"filename": file.filename, "file_id": file_id, "chunks": len(ids)}
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    for f in files:
        summary = await process_one(f)
        total_chunks += summary["chunks"]
        file_summaries.append(summary)

    return {"message": "Ingestion complete", "total_chunks": total_chunks, "files": file_summaries}


class IngestURLRequest(BaseModel):
    url: str


@app.post("/ingest_url")
async def ingest_url(body: IngestURLRequest):
    tmp_path = await run_in_threadpool(download_to_temp, body.url)
    filename = body.url.split("/")[-1].split("?")[0] or "downloaded_file"
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext == ".pdf":
            chunks = await run_in_threadpool(ingest_pdf_streaming_to_store, tmp_path, filename, app.state.vstore, 64, None)
            return {"message": "Ingestion complete", "chunks": chunks, "source": body.url}
        else:
            parsed = await run_in_threadpool(parse_any, tmp_path, filename)
            if not parsed:
                return {"message": "No text extracted", "chunks": 0, "source": body.url}
            texts = [t for t, _ in parsed]
            metas = [m for _, m in parsed]
            ids = await run_in_threadpool(app.state.vstore.add_texts, texts, metas)
            return {"message": "Ingestion complete", "chunks": len(ids), "source": body.url}
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


@app.post("/query")
async def query(req: QueryRequest = Body(...)):
    if app.state.vstore.index.ntotal == 0:
        raise HTTPException(status_code=400, detail="No documents indexed yet. Upload files first.")

    question = req.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question is empty.")

    # Optional OCR on image input
    img_text = ""
    if req.image_base64:
        try:
            img = decode_base64_image(req.image_base64)
            img_text = await run_in_threadpool(image_to_text, img)
        except HTTPException:
            raise
        except Exception as e:
            logger.warning(f"Failed to OCR image in query: {e}")

    augmented_q = question if not img_text else f"{question}\n\nAdditional image text:\n{img_text}"

    # Vector search
    top_k = min(req.top_k or TOP_K_DEFAULT, 5)
    results = await run_in_threadpool(app.state.vstore.search, augmented_q, top_k)
    if not results:
        raise HTTPException(status_code=404, detail="No relevant context found.")

    # Build chat messages and call LLM
    messages = build_messages(question=question, contexts=results)
    answer = await run_in_threadpool(app.state.llm.generate, messages)

    # Enrich results with icons and labels for response
    def enrich(r: Dict[str, Any]) -> Dict[str, Any]:
        icon, label = file_icon_and_label(r)
        return {
            "id": r.get("id"),
            "file_id": r.get("file_id"),
            "filename": r.get("filename"),
            "type": (r.get("type") or "unknown"),
            "type_label": label,
            "icon": icon,
            "page": r.get("page"),
            "table": r.get("table"),
            "chunk_index": r.get("chunk_index"),
            "score": r.get("score"),
            "text": r.get("text"),
        }

    enriched = [enrich(r) for r in results]

    sources = [
        {k: v for k, v in e.items() if k in ["id", "file_id", "filename", "type", "type_label", "icon", "page", "table", "chunk_index", "score"]}
        for e in enriched
    ]
    contexts = [
        {k: v for k, v in e.items() if k in ["text", "file_id", "filename", "type", "type_label", "icon", "page", "chunk_index", "id", "score", "table"]}
        for e in enriched
    ]

    return JSONResponse(content={
        "answer": (answer or "").strip(),
        "context": contexts,
        "sources": sources
    })